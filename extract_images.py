"""
Extract images from HTML files and create a Word document with all articles.

Output structure:
- Name of Article
- URL of article
- Images (embedded in document)

Usage:
    python extract_images.py
    python extract_images.py --output my_images.docx
    python extract_images.py --download  # Download and embed actual images
"""

import os
import re
import argparse
from pathlib import Path
from bs4 import BeautifulSoup
from urllib.parse import unquote
import requests
from io import BytesIO
from concurrent.futures import ThreadPoolExecutor, as_completed

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("python-docx not installed. Installing...")
    import subprocess
    subprocess.check_call(["pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH


def get_article_title_from_filename(filename):
    """Extract a readable title from the filename."""
    # Remove date prefix and extension: 2026-01-27_26-dividend-stocks-for-2026.html -> 26 dividend stocks for 2026
    name = Path(filename).stem
    # Remove date prefix if present (YYYY-MM-DD_)
    if re.match(r'^\d{4}-\d{2}-\d{2}_', name):
        name = name[11:]  # Skip the date and underscore
    # Replace hyphens with spaces and title case
    title = name.replace('-', ' ').replace('_', ' ')
    return title.title()


def get_article_url_from_md(md_path):
    """Extract the source URL from the corresponding markdown file."""
    if not os.path.exists(md_path):
        return None
    
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Find the last Source: line that starts with http
    # Looking for pattern: Source: https://...
    lines = content.strip().split('\n')
    for line in reversed(lines):
        line = line.strip()
        if line.startswith('Source:'):
            url = line.replace('Source:', '').strip()
            if url.startswith('http'):
                return url
    return None


def extract_images_from_html(html_path):
    """Extract all unique image URLs from an HTML file."""
    with open(html_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    soup = BeautifulSoup(content, 'lxml')
    images = set()
    
    # Find all img tags
    for img in soup.find_all('img'):
        src = img.get('src')
        if src and 'substackcdn.com' in src:
            # Clean up the URL - get the highest quality version
            images.add(src)
    
    # Also check for image links in anchor tags (often contain full-res images)
    for a in soup.find_all('a', class_='image-link'):
        href = a.get('href')
        if href and 'substackcdn.com' in href:
            images.add(href)
    
    # Also check picture/source tags
    for source in soup.find_all('source'):
        srcset = source.get('srcset')
        if srcset and 'substackcdn.com' in srcset:
            # srcset contains multiple URLs with sizes, get the largest
            urls = srcset.split(',')
            for url_part in urls:
                url = url_part.strip().split(' ')[0]
                if 'substackcdn.com' in url:
                    images.add(url)
    
    return list(images)


def clean_image_url(url):
    """Clean and decode image URL for better display."""
    # Decode URL-encoded characters
    decoded = unquote(url)
    return decoded


def get_best_image_url(url):
    """Get the best quality version of the image URL."""
    # Substack images often have size/quality parameters
    # Try to get a reasonable size version
    if 'substackcdn.com' in url:
        # Remove size constraints if present for better quality
        # But keep some limit to avoid huge downloads
        url = re.sub(r',w_\d+', ',w_800', url)
        url = re.sub(r',c_limit', ',c_limit', url)
    return url


def download_image(url, timeout=15):
    """Download an image and return as bytes."""
    try:
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        }
        best_url = get_best_image_url(url)
        response = requests.get(best_url, timeout=timeout, headers=headers)
        if response.status_code == 200:
            return url, BytesIO(response.content)
    except Exception as e:
        pass  # Silent fail, will show URL instead
    return url, None


def download_all_images(articles, max_workers=20):
    """Download all images in parallel. Returns dict mapping URL to image data."""
    all_urls = []
    for article in articles:
        all_urls.extend(article['images'])
    
    # Deduplicate URLs
    unique_urls = list(set(all_urls))
    print(f"Downloading {len(unique_urls)} unique images with {max_workers} workers...")
    
    image_cache = {}
    completed = 0
    
    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        futures = {executor.submit(download_image, url): url for url in unique_urls}
        for future in as_completed(futures):
            url, data = future.result()
            image_cache[url] = data
            completed += 1
            if completed % 50 == 0 or completed == len(unique_urls):
                print(f"  Downloaded {completed}/{len(unique_urls)} images...")
    
    success = sum(1 for v in image_cache.values() if v is not None)
    print(f"  Successfully downloaded {success}/{len(unique_urls)} images")
    return image_cache


def create_word_document(articles, output_path, download_images=False, image_cache=None):
    """Create a Word document with all articles and their images."""
    doc = Document()
    
    # Add title
    title = doc.add_heading('Substack Articles - Image Extraction', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add generation info
    from datetime import datetime
    info = doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()  # Spacer
    
    if image_cache is None:
        image_cache = {}
    
    # Process each article
    for i, article in enumerate(articles, 1):
        print(f"Processing article {i}/{len(articles)}: {article['title'][:50]}...")
        
        # Article title
        doc.add_heading(article['title'], level=1)
        
        # Substack/Source info
        source_para = doc.add_paragraph()
        source_para.add_run('Source: ').bold = True
        source_para.add_run(article['substack'])
        
        # Article URL
        url_para = doc.add_paragraph()
        url_para.add_run('URL: ').bold = True
        if article['url']:
            url_para.add_run(article['url'])
        else:
            url_para.add_run('(URL not found)')
        
        # Images section
        doc.add_paragraph()
        images_heading = doc.add_paragraph()
        images_heading.add_run(f'Images ({len(article["images"])} found):').bold = True
        
        if article['images']:
            for j, img_url in enumerate(article['images'], 1):
                if download_images:
                    # Use pre-downloaded image from cache
                    img_data = image_cache.get(img_url)
                    if img_data:
                        try:
                            # Reset stream position
                            img_data.seek(0)
                            doc.add_picture(img_data, width=Inches(5.5))
                            # Add caption
                            caption = doc.add_paragraph(f'Image {j}')
                            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            caption_run = caption.runs[0]
                            caption_run.font.size = Pt(9)
                            caption_run.font.italic = True
                        except Exception as e:
                            # Fall back to URL only
                            img_para = doc.add_paragraph()
                            img_para.add_run(f'{j}. ').bold = True
                            display_url = clean_image_url(img_url)
                            img_para.add_run(display_url[:150] + '...' if len(display_url) > 150 else display_url)
                    else:
                        # URL only if download failed
                        img_para = doc.add_paragraph()
                        img_para.add_run(f'{j}. ').bold = True
                        display_url = clean_image_url(img_url)
                        img_para.add_run(display_url[:150] + '...' if len(display_url) > 150 else display_url)
                else:
                    # Just add the URL
                    img_para = doc.add_paragraph()
                    img_para.add_run(f'{j}. ').bold = True
                    display_url = clean_image_url(img_url)
                    if len(display_url) > 150:
                        display_url = display_url[:150] + '...'
                    img_para.add_run(display_url)
        else:
            doc.add_paragraph('No images found in this article.')
        
        # Add page break between articles (except for last one)
        if i < len(articles):
            doc.add_page_break()
    
    # Save document
    doc.save(output_path)
    print(f"\nDocument saved to: {output_path}")


def main():
    parser = argparse.ArgumentParser(description='Extract images from Substack HTML files to Word document')
    parser.add_argument('--output', type=str, default='substack_images.docx', 
                        help='Output Word document filename (default: substack_images.docx)')
    parser.add_argument('--download', action='store_true',
                        help='Download and embed actual images in parallel')
    parser.add_argument('--workers', type=int, default=20,
                        help='Number of parallel download workers (default: 20)')
    parser.add_argument('--html-dir', type=str, default='html_files',
                        help='Directory containing HTML files (default: html_files)')
    parser.add_argument('--md-dir', type=str, default='md_files',
                        help='Directory containing MD files for URLs (default: md_files)')
    args = parser.parse_args()
    
    html_base = args.html_dir
    md_base = args.md_dir
    
    if not os.path.exists(html_base):
        print(f"Error: HTML directory '{html_base}' not found.")
        return
    
    articles = []
    
    # Scan all substack folders
    for substack_name in sorted(os.listdir(html_base)):
        substack_html_dir = os.path.join(html_base, substack_name)
        substack_md_dir = os.path.join(md_base, substack_name)
        
        if not os.path.isdir(substack_html_dir):
            continue
        
        print(f"\nScanning {substack_name}...")
        
        # Process each HTML file
        for html_file in sorted(os.listdir(substack_html_dir)):
            if not html_file.endswith('.html'):
                continue
            
            html_path = os.path.join(substack_html_dir, html_file)
            md_file = html_file.replace('.html', '.md')
            md_path = os.path.join(substack_md_dir, md_file)
            
            # Extract data
            title = get_article_title_from_filename(html_file)
            url = get_article_url_from_md(md_path)
            images = extract_images_from_html(html_path)
            
            print(f"  Found {len(images)} images in: {title[:50]}...")
            
            articles.append({
                'title': title,
                'substack': substack_name,
                'url': url,
                'images': images,
                'html_file': html_path,
                'md_file': md_path
            })
    
    if not articles:
        print("No articles found!")
        return
    
    print(f"\n{'='*50}")
    print(f"Total: {len(articles)} articles found")
    total_images = sum(len(a['images']) for a in articles)
    print(f"Total images: {total_images}")
    print(f"{'='*50}")
    
    # Pre-download images in parallel if requested
    image_cache = None
    if args.download:
        print(f"\nDownloading images in parallel...")
        image_cache = download_all_images(articles, max_workers=args.workers)
    
    # Create Word document
    print(f"\nCreating Word document...")
    create_word_document(articles, args.output, download_images=args.download, image_cache=image_cache)
    
    print("\nDone!")


if __name__ == "__main__":
    main()
